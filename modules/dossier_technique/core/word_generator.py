import os
import math
from io import BytesIO
# pyrefly: ignore [missing-import]
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT, WD_SECTION

class WordGenerator:
    def __init__(self, output_dir="."):
        self.output_dir = output_dir

    def _render_plot(self, bornes, voisins=None, zoom_out=False, show_grid_ticks=True):
        if not bornes:
            return None
            
        fig, ax = plt.subplots(figsize=(8, 6) if not zoom_out else (4, 4))
        
        # Voisins
        if voisins:
            for nom_voisin, pts_voisin in voisins.items():
                if len(pts_voisin) >= 3:
                    vx = [p[0] for p in pts_voisin]
                    vy = [p[1] for p in pts_voisin]
                    vx.append(vx[0])
                    vy.append(vy[0])
                    ax.plot(vx, vy, 'k--', linewidth=0.5, alpha=0.6)
                    cx = sum(vx[:-1]) / len(pts_voisin)
                    cy = sum(vy[:-1]) / len(pts_voisin)
                    ax.text(cx, cy, nom_voisin, fontsize=7 if zoom_out else 9, 
                            ha='center', va='center', alpha=0.8, fontweight='bold')
                            
        # Lot principal
        xs = [b[0] for b in bornes]
        ys = [b[1] for b in bornes]
        xs.append(xs[0])
        ys.append(ys[0])
        
        ax.plot(xs, ys, 'k-', linewidth=2)
        
        if zoom_out:
            ax.fill(xs, ys, 'k')
            
        if not zoom_out:
            min_x, max_x = min(xs), max(xs)
            min_y, max_y = min(ys), max(ys)
            width = max_x - min_x
            height = max_y - min_y
            margin_x = max(width * 0.3, 10)
            margin_y = max(height * 0.3, 10)
            ax.set_xlim(min_x - margin_x, max_x + margin_x)
            ax.set_ylim(min_y - margin_y, max_y + margin_y)
        
        if not zoom_out:
            for i, (x, y) in enumerate(bornes):
                ax.plot(x, y, 'ko', markersize=3)
                ax.text(x, y, f' B{i+1}', fontsize=10, verticalalignment='bottom')
                
                if i < len(bornes):
                    p1 = bornes[i]
                    p2 = bornes[(i+1)%len(bornes)]
                    dist = math.sqrt((p2[0]-p1[0])**2 + (p2[1]-p1[1])**2)
                    mid_x = (p1[0] + p2[0]) / 2
                    mid_y = (p1[1] + p2[1]) / 2
                    angle = math.degrees(math.atan2(p2[1]-p1[1], p2[0]-p1[0]))
                    if angle > 90: angle -= 180
                    elif angle < -90: angle += 180
                    ax.text(mid_x, mid_y, f"{dist:.3f}", fontsize=8, 
                            ha='center', va='bottom', rotation=angle)
        
        ax.set_aspect('equal')
        
        if show_grid_ticks:
            ax.grid(False)
            ax.set_xticks(ax.get_xticks(), minor=False)
            ax.set_yticks(ax.get_yticks(), minor=False)
            ax.set_xticklabels([])
            ax.set_yticklabels([])
            ax.tick_params(axis='both', which='major', length=0)
            for xt in ax.get_xticks():
                for yt in ax.get_yticks():
                    ax.plot(xt, yt, marker='+', color='grey', markersize=8, alpha=0.5)
        else:
            ax.grid(True, linestyle='--', alpha=0.5)
            ax.set_xticklabels([])
            ax.set_yticklabels([])
            ax.tick_params(axis='both', length=0)
            
        if zoom_out:
            ax.annotate('N', xy=(0.9, 0.9), xycoords='axes fraction', 
                        xytext=(0.9, 0.75), textcoords='axes fraction',
                        arrowprops=dict(facecolor='black', width=2, headwidth=8),
                        fontsize=12, ha='center', va='top')
                        
        for spine in ax.spines.values():
            spine.set_linewidth(1)
            
        plt.tight_layout()
        
        buf = BytesIO()
        plt.savefig(buf, format='png', dpi=200, transparent=True, bbox_inches='tight', pad_inches=0.1)
        plt.close(fig)
        buf.seek(0)
        return buf

    def _format_surface_text(self, surface_m2):
        ha = int(surface_m2 // 10000)
        a = int((surface_m2 % 10000) // 100)
        ca = int(surface_m2 % 100)
        return f"{ha:02d} Ha {a:02d} A {ca:02d} Ca"

    def generate_word(self, filename, data):
        def v(val): return val if val and str(val).strip() else "......"
            
        doc = Document()
        
        # === Styles de base ===
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(12)
        
        # =======================
        # PAGE 1 : GARDE
        # =======================
        table = doc.add_table(rows=1, cols=2)
        table.autofit = True
        cell1 = table.cell(0, 0)
        cell1.text = "MINISTERE DES FINANCES ET DU BUDGET\nDIRECTION GENERALE DES IMPOTS\nDIRECTION DU CADASTRE"
        
        cell2 = table.cell(0, 1)
        cell2.text = "REPUBLIQUE DE COTE D'IVOIRE\nUnion – Discipline - Travail"
        cell2.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        doc.add_paragraph("\n\n")
        
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("DOSSIER DE CALCULS")
        run.bold = True
        run.font.size = Pt(36)
        
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("DOSSIER TECHNIQUE DE MORCELLEMENT")
        run.bold = True
        run.underline = True
        run.font.size = Pt(20)
        
        doc.add_paragraph("\n\n")
        
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"CENTRE : {v(data.get('centre'))}\n")
        run.bold = True
        run.font.size = Pt(14)
        run = p.add_run(f"Ilot : {v(data.get('ilot'))}          Lot : {v(data.get('lot'))}\n")
        run.bold = True
        run.font.size = Pt(14)
        run = p.add_run(f"Demandeur : {v(data.get('demandeur'))}")
        run.bold = True
        run.font.size = Pt(14)
        
        doc.add_page_break()
        
        # =======================
        # PAGE 2 : RAPPORT
        # =======================
        table = doc.add_table(rows=1, cols=2)
        cell1 = table.cell(0, 0)
        cell1.text = "MINISTERE DES FINANCES ET DU BUDGET\nDIRECTION DES IMPOTS\nDIRECTION DU CADASTRE"
        cell2 = table.cell(0, 1)
        cell2.text = "REPUBLIQUE DE COTE D'IVOIRE"
        cell2.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        doc.add_paragraph("\n")
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("RAPPORT DU GEOMETRE")
        run.bold = True
        run.underline = True
        run.font.size = Pt(18)
        
        table = doc.add_table(rows=3, cols=2)
        table.cell(0,0).text = f"Livre Foncier : {v(data.get('livre_foncier'))}"
        table.cell(0,1).text = f"Centre : {v(data.get('centre'))}"
        table.cell(1,0).text = f"Lotissement : {v(data.get('lotissement'))}"
        table.cell(1,1).text = f"Ilot : {v(data.get('ilot'))}    Lot : {v(data.get('lot'))}"
        table.cell(2,0).text = f"Morcellement de TF : {v(data.get('tf'))}"
        table.cell(2,1).text = f"Section : {v(data.get('section'))}"
        
        doc.add_paragraph(f"\nDate de bornage fait sur le terrain par le Géomètre : Antérieure\nDate de consultation des documents cadastraux : {v(data.get('date_consultation'))}\n")
        
        p = doc.add_paragraph()
        p.add_run("DOCUMENTS DE BASE UTILISES\n").bold = True
        p.add_run(f"CADASTRE :\nN° de la section du plan : {v(data.get('section'))}\nN° du dossier : {v(data.get('dossier'))}\n")
        p.add_run(f"GEOMETRE PRIVE :\nNom du cabinet : {v(data.get('cabinet_nom'))}")
        
        p = doc.add_paragraph("LES COORDONNEES DOIVENT ETRE CELLES DU SYSTEME WGS 84 UTM FUSEAU 30N\nCOORDONNEES DES SOMMETS DE L'ILOT UTILISE OU POLYGONATION")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Tableau Coordonnees
        bornes = data.get("bornes", [])
        if bornes:
            table = doc.add_table(rows=len(bornes)+1, cols=3)
            table.style = 'Table Grid'
            table.autofit = False
            table.allow_autofit = False
            
            hdr = table.rows[0].cells
            hdr[0].text, hdr[1].text, hdr[2].text = "Bornes", "X", "Y"
            for i, borne in enumerate(bornes):
                row = table.rows[i+1].cells
                row[0].text = f"B{i+1}"
                row[1].text = f"{borne[0]:.3f}"
                row[2].text = f"{borne[1]:.3f}"
                
            # Ajustement des largeurs
            col_widths = [Cm(2.0), Cm(4.0), Cm(4.0)]
            for row in table.rows:
                row.height = Cm(0.8)
                for idx, width in enumerate(col_widths):
                    row.cells[idx].width = width
                
        # =======================
        # PAGE 3 : CALCUL RETOUR
        # =======================
        doc.add_page_break()
        
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("TABLEAU DE COORDONNEES\n")
        run.bold = True
        run.font.size = Pt(14)
        run = p.add_run("CALCUL RETOUR ET CALCUL DE SURFACE")
        run.italic = True
        run.bold = True
        run.font.size = Pt(12)
        
        doc.add_paragraph("\n")
        
        # Generation de bornes_calc indentique au PDF
        bornes_calc = []
        if bornes:
            for i in range(len(bornes)):
                b1 = bornes[i]
                bornes_calc.append({"point": f"P{i+1}", "nom": f"B{i+1}", "x": b1[0], "y": b1[1], "angle": 100.0, "dist": None, "gis": None})
            
            for i in range(len(bornes)):
                b1 = bornes[i]
                b2 = bornes[(i+1)%len(bornes)]
                dx = b2[0] - b1[0]
                dy = b2[1] - b1[1]
                dist = math.sqrt(dx**2 + dy**2)
                gis = math.atan2(dx, dy) * 200 / math.pi
                if gis < 0: gis += 400
                
                if i == len(bornes) - 1:
                    bornes_calc.append({"point": "P1", "nom": "B1", "x": None, "y": None, "angle": None, "dist": dist, "gis": gis})
                else:
                    bornes_calc[i+1]["dist"] = dist
                    bornes_calc[i+1]["gis"] = gis
            
            ctable = doc.add_table(rows=len(bornes_calc)+2, cols=7)
            ctable.style = 'Table Grid'
            ctable.autofit = False
            ctable.allow_autofit = False
            
            hdr = ctable.rows[0].cells
            hdr[0].merge(hdr[3]).text = "COORDONNEES"
            hdr[4].merge(hdr[6]).text = "CALCUL RETOUR"
            
            sub = ctable.rows[1].cells
            sub[0].text, sub[1].text, sub[2].text, sub[3].text, sub[4].text, sub[5].text, sub[6].text = "POINT", "BORNE", "X", "Y", "ANGLES", "DISTANCES", "GISEMENTS"
            
            for i, b in enumerate(bornes_calc):
                row = ctable.rows[i+2].cells
                row[0].text = b["point"] if b["point"] else ""
                row[1].text = b["nom"]
                row[2].text = f"{b['x']:.3f}" if b['x'] is not None else ""
                row[3].text = f"{b['y']:.3f}" if b['y'] is not None else ""
                row[4].text = f"{b['angle']:.3f}" if b['angle'] is not None else ""
                row[5].text = f"{b['dist']:.3f}" if b['dist'] is not None else ""
                row[6].text = f"{b['gis']:.3f}" if b['gis'] is not None else ""

            # Ajustement des largeurs de colonnes et hauteurs de lignes
            col_widths = [Cm(1.5), Cm(1.7), Cm(3.0), Cm(3.0), Cm(2.0), Cm(2.5), Cm(2.5)]
            for row in ctable.rows:
                row.height = Cm(0.8)  # Augmente la hauteur de la ligne
                for idx, width in enumerate(col_widths):
                    row.cells[idx].width = width
                
        # =======================
        # PAGE 4 : EXTRAIT (PAYSAGE)
        # =======================
        # La page 4 a été retirée selon la demande.
                                    
        output_path = os.path.join(self.output_dir, filename)
        doc.save(output_path)
        return output_path
